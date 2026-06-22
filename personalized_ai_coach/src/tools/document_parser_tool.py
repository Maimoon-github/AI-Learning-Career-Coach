from __future__ import annotations

import asyncio
import os
import json
from pathlib import Path
from typing import Any, List, Optional, Dict, Union
from datetime import datetime

import structlog
from crewai.tools import BaseTool
from pydantic import BaseModel, Field

# Local imports
try:
    from src.utils.error_handling import ToolError
except ImportError:
    class ToolError(Exception): pass

log = structlog.get_logger(__name__)

# ---------------------------------------------------------------------------
# Output Models
# ---------------------------------------------------------------------------

class ParsingMetadata(BaseModel):
    """Rich metadata for a parsed document."""
    source_path: str
    extension: str
    size_bytes: int
    created_at: Optional[str] = None
    modified_at: Optional[str] = None
    page_count: Optional[int] = None
    table_count: int = 0
    image_count: int = 0
    parsing_confidence: float = 1.0
    parser_used: str
    timestamp: str = Field(default_factory=lambda: datetime.now().isoformat())

class ParsingResult(BaseModel):
    """Structured output for a single document."""
    success: bool
    text: str = ""
    markdown: str = ""
    tables: List[Dict[str, Any]] = []
    metadata: ParsingMetadata
    error: Optional[str] = None

# ---------------------------------------------------------------------------
# Input Schema
# ---------------------------------------------------------------------------

class DocumentParserInput(BaseModel):
    """Input for parsing one or more documents."""
    path: str = Field(
        description="Absolute path to a file or directory. Supports PDF, DOCX, MD, TXT, etc."
    )
    recursive: bool = Field(
        default=False, 
        description="If path is a directory, search recursively for files."
    )
    use_llamaparse: bool = Field(
        default=False,
        description="Force use of LlamaParse (requires LLAMA_CLOUD_API_KEY). Best for complex layouts."
    )
    max_files: int = Field(
        default=10,
        description="Maximum number of files to parse in a directory."
    )

# ---------------------------------------------------------------------------
# Tool Implementation
# ---------------------------------------------------------------------------

class DocumentParserTool(BaseTool):
    """
    Advanced Document Parser using modern layout-aware engines.
    
    Sequence:
    1. Docling (Primary for PDF/Layouts)
    2. MarkItDown (Best for MS Office/Miscellaneous)
    3. Legacy (PyPDF, Docx) as baseline
    4. LlamaParse (Cloud fallback/optional high-quality)
    """
    name: str = "document_parser"
    description: str = (
        "Parses documents (PDF, DOCX, MD, TXT) into high-quality Markdown. "
        "Extracts text, tables, and rich metadata. Works on single files or directories."
    )
    args_schema: type[BaseModel] = DocumentParserInput

    # ------------------------------------------------------------------
    # Public Entry Points
    # ------------------------------------------------------------------

    def _run(
        self, 
        path: str, 
        recursive: bool = False, 
        use_llamaparse: bool = False,
        max_files: int = 10
    ) -> str:
        """Sync entry point for CrewAI."""
        try:
            loop = asyncio.get_event_loop()
        except RuntimeError:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)

        if loop.is_running():
            # Bridge to async run if we are in an existing loop
            import threading
            from concurrent.futures import Future

            def _run_in_new_loop(fut: Future):
                new_loop = asyncio.new_event_loop()
                asyncio.set_event_loop(new_loop)
                try:
                    result = new_loop.run_until_complete(
                        self._async_run(path, recursive, use_llamaparse, max_files)
                    )
                    fut.set_result(result)
                except Exception as e:
                    fut.set_exception(e)
                finally:
                    new_loop.close()

            f: Future = Future()
            t = threading.Thread(target=_run_in_new_loop, args=(f,))
            t.start()
            t.join()
            return f.result()
        else:
            return loop.run_until_complete(
                self._async_run(path, recursive, use_llamaparse, max_files)
            )

    async def _async_run(
        self, 
        path: str, 
        recursive: bool = False, 
        use_llamaparse: bool = False,
        max_files: int = 10
    ) -> str:
        """Core async logic."""
        p = Path(path)
        if not p.exists():
            return f"Error: Path not found: {path}"

        if p.is_dir():
            results = await self.parse_directory(p, recursive, use_llamaparse, max_files)
            return self._format_batch_output(results)
        else:
            result = await self.parse_file(p, use_llamaparse)
            return self._format_single_output(result)

    # ------------------------------------------------------------------
    # Core Logic
    # ------------------------------------------------------------------

    async def parse_file(self, file_path: Path, use_llamaparse: bool = False) -> ParsingResult:
        """Parse a single file using the fallback chain."""
        log.info("parsing_file_start", path=str(file_path), use_llamaparse=use_llamaparse)
        
        # 0. Initial Metadata
        stats = file_path.stat()
        meta = ParsingMetadata(
            source_path=str(file_path),
            extension=file_path.suffix.lower(),
            size_bytes=stats.st_size,
            created_at=datetime.fromtimestamp(stats.st_ctime).isoformat(),
            modified_at=datetime.fromtimestamp(stats.st_mtime).isoformat(),
            parser_used="none"
        )

        # 1. Option: LlamaParse (Explicit or if configured)
        if use_llamaparse or os.getenv("LLAMA_CLOUD_API_KEY"):
            # If use_llamaparse is False but we have the key, we might still want it as a last resort
            # but for now we follow the user's "optional support" / "explicit" hint.
            if use_llamaparse:
                res = await self._try_llamaparse(file_path, meta)
                if res.success: return res

        # 2. Primary: Docling
        res = await self._try_docling(file_path, meta)
        if res.success: return res

        # 3. Secondary: MarkItDown
        res = await self._try_markitdown(file_path, meta)
        if res.success: return res

        # 4. Tertiary: Legacy Parsers
        res = await self._try_legacy(file_path, meta)
        if res.success: return res

        # 5. Quaternary: Unstructured (for niche formats)
        res = await self._try_unstructured(file_path, meta)
        if res.success: return res

        # 6. Last Resort: Plain Text
        return await self._try_plain_text(file_path, meta)

    async def parse_directory(
        self, 
        dir_path: Path, 
        recursive: bool, 
        use_llamaparse: bool,
        max_files: int
    ) -> List[ParsingResult]:
        """Parse multiple files in a directory."""
        pattern = "**/*" if recursive else "*"
        files = [f for f in dir_path.glob(pattern) if f.is_file()][:max_files]
        
        tasks = [self.parse_file(f, use_llamaparse) for f in files]
        return await asyncio.gather(*tasks)

    # ------------------------------------------------------------------
    # Specific Parsers
    # ------------------------------------------------------------------

    async def _try_docling(self, path: Path, meta: ParsingMetadata) -> ParsingResult:
        """Use Docling for high-quality document conversion."""
        try:
            from docling.document_converter import DocumentConverter
            
            # Docling works well for PDF, DOCX, HTML, MD
            supported = {".pdf", ".docx", ".html", ".md", ".pptx"}
            if path.suffix.lower() not in supported:
                return ParsingResult(success=False, metadata=meta)

            converter = DocumentConverter()
            # run_in_executor since conversion can be CPU intensive and blocking
            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(None, converter.convert, str(path))
            
            doc = result.document
            markdown = doc.export_to_markdown()
            
            # Extract tables if available
            tables = []
            # Note: Docling's table extraction API might vary, 
            # this is a placeholder for actual table objects in their schema.
            # Usually tables are in the markdown, but structured data is better.
            
            meta.parser_used = "docling"
            meta.parsing_confidence = 0.95
            # meta.page_count = len(doc.pages) if hasattr(doc, 'pages') else None
            
            return ParsingResult(
                success=True,
                markdown=markdown,
                text=doc.export_to_text(),
                tables=tables,
                metadata=meta
            )
        except Exception as e:
            log.debug("docling_failed", path=str(path), error=str(e))
            return ParsingResult(success=False, metadata=meta, error=str(e))

    async def _try_markitdown(self, path: Path, meta: ParsingMetadata) -> ParsingResult:
        """Microsoft MarkItDown for various office formats."""
        try:
            from markitdown import MarkItDown
            
            md = MarkItDown()
            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(None, md.convert, str(path))
            
            meta.parser_used = "markitdown"
            meta.parsing_confidence = 0.85
            
            return ParsingResult(
                success=True,
                markdown=result.text_content,
                text=result.text_content, # MarkItDown is text-centric
                metadata=meta
            )
        except Exception as e:
            log.debug("markitdown_failed", path=str(path), error=str(e))
            return ParsingResult(success=False, metadata=meta, error=str(e))

    async def _try_llamaparse(self, path: Path, meta: ParsingMetadata) -> ParsingResult:
        """Cloud-based LlamaParse for complex documents."""
        api_key = os.getenv("LLAMA_CLOUD_API_KEY")
        if not api_key:
            return ParsingResult(success=False, metadata=meta, error="No LLAMA_CLOUD_API_KEY")

        try:
            from llama_parse import LlamaParse
            
            parser = LlamaParse(result_type="markdown", api_key=api_key)
            # LlamaParse often has async methods
            json_results = await parser.aget_json(str(path))
            
            if not json_results:
                return ParsingResult(success=False, metadata=meta)

            # LlamaParse returns list of dicts (one per file, though we passing one)
            data = json_results[0]
            markdown = data.get("markdown", "")
            
            meta.parser_used = "llamaparse"
            meta.parsing_confidence = 0.98
            meta.page_count = data.get("page_count")
            
            return ParsingResult(
                success=True,
                markdown=markdown,
                text=markdown, # Best representation is the markdown itself
                metadata=meta
            )
        except Exception as e:
            log.debug("llamaparse_failed", path=str(path), error=str(e))
            return ParsingResult(success=False, metadata=meta, error=str(e))

    async def _try_legacy(self, path: Path, meta: ParsingMetadata) -> ParsingResult:
        """Fallback to PyPDF and python-docx."""
        suffix = path.suffix.lower()
        try:
            if suffix == ".pdf":
                from pypdf import PdfReader
                reader = PdfReader(str(path))
                text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
                meta.parser_used = "pypdf"
                meta.page_count = len(reader.pages)
                return ParsingResult(success=True, text=text, markdown=text, metadata=meta)
                
            elif suffix == ".docx":
                from docx import Document
                doc = Document(str(path))
                text = "\n".join(p.text for p in doc.paragraphs)
                meta.parser_used = "python-docx"
                return ParsingResult(success=True, text=text, markdown=text, metadata=meta)
            
            return ParsingResult(success=False, metadata=meta)
        except Exception as e:
            log.debug("legacy_failed", path=str(path), error=str(e))
            return ParsingResult(success=False, metadata=meta, error=str(e))

    async def _try_unstructured(self, path: Path, meta: ParsingMetadata) -> ParsingResult:
        """Use Unstructured for broad format support."""
        try:
            from unstructured.partition.auto import partition
            
            loop = asyncio.get_event_loop()
            elements = await loop.run_in_executor(None, partition, str(path))
            text = "\n\n".join([str(el) for el in elements])
            
            meta.parser_used = "unstructured"
            meta.parsing_confidence = 0.70
            
            return ParsingResult(
                success=True,
                text=text,
                markdown=text,
                metadata=meta
            )
        except Exception as e:
            log.debug("unstructured_failed", path=str(path), error=str(e))
            return ParsingResult(success=False, metadata=meta, error=str(e))

    async def _try_plain_text(self, path: Path, meta: ParsingMetadata) -> ParsingResult:
        """Last resort: read as plain text."""
        try:
            content = path.read_text(encoding="utf-8", errors="replace")
            meta.parser_used = "plain_text"
            return ParsingResult(success=True, text=content, markdown=content, metadata=meta)
        except Exception as e:
            return ParsingResult(success=False, metadata=meta, error=str(e))

    # ------------------------------------------------------------------
    # Formatting
    # ------------------------------------------------------------------

    def _format_single_output(self, result: ParsingResult) -> str:
        """Format a single parsing result into a pretty markdown string."""
        if not result.success:
            return f"### Failed to parse: {result.metadata.source_path}\nError: {result.error}"

        m = result.metadata
        header = (
            f"## Document: {Path(m.source_path).name}\n"
            f"- **Parser**: {m.parser_used}\n"
            f"- **Confidence**: {m.parsing_confidence:.2f}\n"
            f"- **Pages**: {m.page_count or 'N/A'}\n"
            f"- **Format**: {m.extension}\n"
            f"---\n"
        )
        
        # Prefer markdown as it contains structure (tables, headings)
        body = result.markdown if result.markdown else result.text
        
        # Add a summary of tables if they were extracted but not in MD (unlikely with docling)
        table_info = ""
        if result.tables:
            table_info = f"\n\n*Extracted {len(result.tables)} tables.*\n"

        return header + body + table_info

    def _format_batch_output(self, results: List[ParsingResult]) -> str:
        """Format multiple results."""
        outputs = [self._format_single_output(r) for r in results]
        return "\n\n" + ("=" * 40) + "\n\n".join(outputs)


# ---------------------------------------------------------------------------
# Batch Helper
# ---------------------------------------------------------------------------

def parse_multiple(file_paths: List[str]) -> List[ParsingResult]:
    """Helper for batch parsing outside the tool context."""
    tool = DocumentParserTool()
    results = []
    # We use a simple loop here, but the tool itself supports directories
    for path in file_paths:
        results.append(asyncio.run(tool.parse_file(Path(path))))
    return results